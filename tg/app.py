import logging
import os
from datetime import datetime
from docx import Document
from openpyxl import Workbook, load_workbook
from telegram import Update, InputFile, ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import (
    ApplicationBuilder, ContextTypes,
    CommandHandler, MessageHandler, filters
)
from dotenv import load_dotenv
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Font, Border, Side, Alignment
import requests

# === Настройки ===
load_dotenv()
BOT_TOKEN = os.getenv("TELEGRAM_TOKEN", "7974253432:AAFztKIqek8xTPotu3NvtXucc26RLFRoFyE")
DEEPSEEK_API_KEY = "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"
DEEPSEEK_API_URL = "https://openrouter.ai/api/v1/chat/completions"

TEMPLATE_FIELDS = [
    "Название",
    "Что хотим улучшить?",
    "Какие данные поступают агенту на выход?",
    "Как процесс выглядит сейчас? as-is",
    "Какой результат нужен от агента?",
    "Достижимый идеал(to-be)",
    "Масштаб процесса"
]

user_states = {}
agent_query_state = {}

# === Проверка идеи через DeepSeek с анализом файла agents.xlsx ===
def check_idea_with_deepseek(user_input: str) -> tuple[str, str]:
    try:
        wb = load_workbook("agents.xlsx")
        ws = wb.active
        all_agents_data = []
        contact = None
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row[0]:
                continue
            name, team, contact_cell, desc = row
            full_info = f"Название: {name}, Команда: {team}, Контакт: {contact_cell}, Описание: {desc}"
            all_agents_data.append(full_info)
            if name and user_input.lower() in name.lower():
                contact = contact_cell
        joined_data = "\n".join(all_agents_data)
    except Exception as e:
        joined_data = "(не удалось загрузить данные об агентах)"
        contact = None

    prompt = (
        f"Вот список существующих AI-агентов:\n{joined_data}\n\n"
        f"Пользователь предлагает идею: {user_input}.\n"
        "Проверь, есть ли похожие идеи и оцени, насколько она уникальна. Ответь кратко и по делу. Если идеа уникальна, полностью и частично, напиши Контакт лидера в конце ответа!"
    )

    headers = {
        'Authorization': f'Bearer {DEEPSEEK_API_KEY}',
        'Content-Type': 'application/json'
    }
    data = {
        "model": "deepseek/deepseek-chat:free",
        "messages": [{"role": "user", "content": prompt}]
    }
    try:
        response = requests.post(DEEPSEEK_API_URL, json=data, headers=headers)
        if response.status_code == 200:
            summary = response.json()['choices'][0]['message']['content']
            return summary, contact
        else:
            return f"Ошибка DeepSeek API: {response.status_code}", None
    except Exception as e:
        return f"⚠️ Ошибка при обращении к DeepSeek: {e}", None

# === Главное меню ===
def get_main_menu():
    keyboard = [
        [KeyboardButton("У меня есть идея!💡")],
        [KeyboardButton("АИ-агенты?📍")],
        [KeyboardButton("Кто поможет?💬")],
        [KeyboardButton("Поддержка📝")]
    ]
    return ReplyKeyboardMarkup(keyboard, resize_keyboard=True, one_time_keyboard=True)

# === Генерация файлов ===
def generate_files(data):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    word_path = f"agent_{timestamp}.docx"
    excel_path = f"agent_{timestamp}.xlsx"

    doc = Document()
    title = doc.add_heading("AI-агент — шаблон", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for key, value in data.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{key}:\n")
        run.bold = True
        run.font.size = Pt(14)

        run2 = p.add_run(f"{value}\n")
        run2.font.size = Pt(12)
        p.space_after = Pt(12)

    doc.save(word_path)

    wb = Workbook()
    ws = wb.active
    ws.title = "Агент"

    bold_font = Font(bold=True)
    thin_border = Border(left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin"))
    alignment = Alignment(wrap_text=True, vertical="top")

    ws.append(["Поле", "Значение"])
    for cell in ws[1]:
        cell.font = bold_font
        cell.border = thin_border
        cell.alignment = alignment

    for key, value in data.items():
        ws.append([key, value])
        for cell in ws[ws.max_row]:
            cell.border = thin_border
            cell.alignment = alignment

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 60

    wb.save(excel_path)
    return word_path, excel_path

# === Команда /start ===
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        f"👋 Привет, @{update.effective_user.username or 'user'}! \n"
        "Меня зовут <b>Агентолог</b>, я помогу тебе с идеями для AI-агентов.\n\n"
        "Вот что я могу сделать:\n"
        "<b>1. У меня есть идея!💡</b>\n"
        "Я помогу тебе узнать, насколько твоя идея <b>уникальна!</b>\n\n"
        "<b>2. АИ-агенты?📍</b>\n"
        "АИ-агенты разрабатываются каждый день, здесь мы собрали самый <b>свежий список агентов!</b>\n\n"
        "<b>3. Кто поможет?💬</b>\n"
        "Агентов очень много и не всегда можно найти, кто их разрабатывает. Давай <b>подскажем</b>, кто эти люди!\n\n"
        "<b>4. Поддержка📝</b>\n"
        "Остались <b>вопросы</b> или <b>предложения</b> по работе чат-бота? Пиши нам!\n\n"
        "Скорее выбирай, что мы будем делать👇",
        reply_markup=get_main_menu(),
        parse_mode="HTML"
    )

# === Обработка сообщений ===
async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    msg = update.message.text.strip()

    if msg == "У меня есть идея!💡":
        user_states[user_id] = {"deepseek_mode": True}
        await update.message.reply_text("💬 Опиши свою идею свободно, я проверю её уникальность:")
        return

    if user_id in user_states and user_states[user_id].get("deepseek_mode"):
        idea_text = msg
        await update.message.reply_text("🔍 Отправляю идею в DeepSeek...")
        result, contact = check_idea_with_deepseek(idea_text)
        await update.message.reply_text(f"🤖 Ответ DeepSeek:\n\n<b>{result}</b>", parse_mode="HTML")

        return

    if msg == "АИ-агенты?📍":
        keyboard = [[KeyboardButton("Все агенты (Excel)")], [KeyboardButton("Искать по названию")]]
        await update.message.reply_text("📋 Что хотите сделать?", reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True))
        return

    if msg == "Все агенты (Excel)":
        try:
            with open("agents.xlsx", "rb") as file:
                await update.message.reply_document(InputFile(file))
        except Exception as e:
            await update.message.reply_text(f"❌ Не удалось отправить файл: {e}")
        return

    if msg == "Искать по названию":
        agent_query_state[user_id] = True
        await update.message.reply_text("🔍 Введи название агента:")
        return

    if agent_query_state.get(user_id):
        agent_query_state[user_id] = False
        term = msg.lower()
        try:
            wb = load_workbook("agents.xlsx")
            ws = wb.active
            results = [r for r in ws.iter_rows(min_row=2, values_only=True) if term in (r[0] or '').lower()]
            if not results:
                await update.message.reply_text("❌ Агент не найден.", reply_markup=get_main_menu())
            else:
                for r in results:
                    name, team, contact, desc = r
                    await update.message.reply_text(
                        f"<b>Название:</b> {name}\n<b>Команда:</b> {team}\n<b>Контакт:</b> {contact}\n<b>Описание:</b> {desc}",
                        parse_mode="HTML")
        except Exception as e:
            await update.message.reply_text(f"⚠️ Ошибка при поиске: {e}", reply_markup=get_main_menu())
        return

    if msg == "Кто поможет?💬":
        await update.message.reply_text("🧑‍💻 Пока эта функция в разработке.", reply_markup=get_main_menu())
        return

    if msg == "Поддержка📝":
        await update.message.reply_text("✉️ Напиши нам в Telegram: @your_support_bot", reply_markup=get_main_menu())
        return

    if user_id in user_states:
        state = user_states[user_id]
        step = state.get("step")

        if step is not None and step < len(TEMPLATE_FIELDS):
            field = TEMPLATE_FIELDS[step]
            state["data"][field] = msg
            state["step"] += 1

            if state["step"] < len(TEMPLATE_FIELDS):
                next_field = TEMPLATE_FIELDS[state["step"]]
                await update.message.reply_text(f"{state['step'] + 1}️⃣ {next_field}:")
            else:
                await update.message.reply_text("✅ Формирую файлы...")
                word_path, excel_path = generate_files(state["data"])

                with open(word_path, "rb") as doc_file:
                    await update.message.reply_document(InputFile(doc_file))
                with open(excel_path, "rb") as excel_file:
                    await update.message.reply_document(InputFile(excel_file))

                os.remove(word_path)
                os.remove(excel_path)
                user_states.pop(user_id)
                await update.message.reply_text("📁 Шаблоны готовы. Выбирай следующее действие:", reply_markup=get_main_menu())
        return

    await update.message.reply_text("🤖 Я вас не понял. Пожалуйста, выбери действие из меню!", reply_markup=get_main_menu())

# === Запуск ===
logging.basicConfig(level=logging.INFO)
app = ApplicationBuilder().token(BOT_TOKEN).build()
app.add_handler(CommandHandler("start", start))
app.add_handler(MessageHandler(filters.TEXT & (~filters.COMMAND), handle_message))

if __name__ == "__main__":
    print("✅ Бот запущен!")
    app.run_polling()
