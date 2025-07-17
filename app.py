from flask import Flask, request, jsonify, send_file
import os
import requests
from docx import Document
from openpyxl import Workbook
from dotenv import load_dotenv
from datetime import datetime
from openpyxl import load_workbook

from init_excel import generate_initiatives_excel

load_dotenv()

app = Flask(__name__)

SBERCHAT_TOKEN = os.getenv("SBERCHAT_TOKEN")
GIGACHAT_API = os.getenv("GIGACHAT_API")
CERT_PATH = os.getenv("CERT_PATH")
KEY_PATH = os.getenv("KEY_PATH")

def load_known_agents_from_excel(path="agents.xlsx"):
    agents = []
    try:
        wb = load_workbook(path)
        ws = wb.active
        for row in ws.iter_rows(min_row=2, max_col=1, values_only=True):
            name = row[0]
            if name:
                agents.append(name.strip().lower())
    except Exception as e:
        print(f"⚠️ Ошибка загрузки Excel: {e}")
    return agents

KNOWN_AGENTS = load_known_agents_from_excel("agents.xlsx")

def is_authorized(req):
    return req.headers.get("Authorization") == f"Bearer {SBERCHAT_TOKEN}"

def generate_files(data):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    word_path = f"agent_{timestamp}.docx"
    excel_path = f"agent_{timestamp}.xlsx"

    doc = Document()
    doc.add_heading("AI-агент (шаблон)", 0)
    for key, value in data.items():
        doc.add_heading(key, level=1)
        doc.add_paragraph(value)
    doc.save(word_path)

    wb = Workbook()
    ws = wb.active
    ws.title = "Агент"
    ws.append(["Поле", "Значение"])
    for key, value in data.items():
        ws.append([key, value])
    wb.save(excel_path)

    return word_path, excel_path

def process_with_gigachat(prompt):
    try:
        response = requests.post(
            GIGACHAT_API,
            json={"prompt": prompt},
            cert=(CERT_PATH, KEY_PATH),
            verify=False
        )
        return response.json().get("response", "Нет ответа от GigaChat.")
    except Exception as e:
        return f"Ошибка при подключении к GigaChat: {str(e)}"

@app.route("/webhook", methods=["POST"])
def webhook():
    if not is_authorized(request):
        return jsonify({"error": "Unauthorized"}), 403

    data = request.json
    user_message = data.get("message", "").strip().lower()

    if user_message == "/start":
        response = "👋 Привет! Я помогу тебе с идеями для AI-агентов.\nНажми кнопку ниже, чтобы начать."
        buttons = ["Старт"]
        return jsonify({
            "reply": response,
            "buttons": buttons
        })
    
    if user_message == "старт":
        response = (
            "Добро пожаловать! Вот что я умею:\n"
            "- Напиши 'шаблон' — чтобы оформить идею\n"
            "- Напиши 'инициатива: ...' — чтобы отправить инициативу\n"
            "- Напиши 'гигачат: твоя идея' — чтобы отправить в GigaChat\n"
        )
        return jsonify({
            "reply": response
        })

    if "шаблон" in user_message:
        template_fields = ["Название", "Целевая аудитория", "Проблема", "Решение", "Каналы", "Технологии"]
        response = "Давайте заполним шаблон. Ответьте на следующие вопросы:\n"
        response += "\n".join([f"{i+1}. {field}" for i, field in enumerate(template_fields)])
        return jsonify({"reply": response})

    elif user_message.startswith("заполнено:"):
        raw = user_message.replace("заполнено:", "").strip()
        pairs = [i.split("=") for i in raw.split(",")]
        data_dict = {k.strip(): v.strip() for k, v in pairs if len(k.strip()) > 0 and len(v.strip()) > 0}
        word_path, excel_path = generate_files(data_dict)
        return jsonify({
            "reply": f"✅ Ваш шаблон оформлен.\n📄 Word: {word_path}\n📊 Excel: {excel_path}"
        })

    elif user_message.startswith("инициатива:"):
        raw = user_message.replace("инициатива:", "").strip()
        pairs = [i.split("=") for i in raw.split(",")]
        initiative = {k.strip().capitalize(): v.strip() for k, v in pairs if len(k.strip()) > 0 and len(v.strip()) > 0}

        agent_name = initiative.get("Название", "").lower()

        if agent_name not in KNOWN_AGENTS:
            return jsonify({
                "reply": "👏 Отличная идея! Давайте оформим её по шаблону. Напишите 'шаблон' для начала."
            })
        else:
            # Сохраняем файл во временную директорию и отправляем
            temp_dir = tempfile.gettempdir()
            temp_path = os.path.join(temp_dir, f"initiative_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
            from init_excel import generate_initiatives_excel
            generate_initiatives_excel([initiative]).replace(temp_path)

            return send_file(
                temp_path,
                as_attachment=True,
                download_name=os.path.basename(temp_path),
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )


    elif "гигачат:" in user_message:
        prompt = user_message.replace("гигачат:", "").strip()
        result = process_with_gigachat(prompt)
        return jsonify({"reply": f"🤖 Ответ GigaChat:\n{result}"})

    else:
        return jsonify({
            "reply": "👋 Привет! Я помогу тебе с идеями для AI-агентов.\n\nНапиши 'шаблон' — чтобы оформить идею, 'инициатива: ...' — чтобы отправить инициативу, или 'гигачат: твоя идея' — чтобы отправить в GigaChat."
        })

if __name__ == "__main__":
    app.run(port=8080)
