import logging
import os
import uuid
import base64
import requests
from datetime import datetime, timedelta
from docx import Document
from openpyxl import Workbook, load_workbook
from dialog_bot_sdk.bot import DialogBot
from dialog_bot_sdk.entities.peers import Peer
from dialog_bot_sdk.utils import AsyncTask
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Font, Border, Side, Alignment
from aiohttp import web

from dotenv import load_dotenv
load_dotenv()

# === Настройки ===
GIGACHAT_AUTH_KEY = os.getenv("GIGACHAT_AUTH_KEY")
GIGACHAT_SCOPE = os.getenv("GIGACHAT_SCOPE")
BOT_TOKEN = os.getenv("BOT_TOKEN")
GIGACHAT_TOKEN_URL = "https://ngw.devices.sberbank.ru:9443/api/v2/oauth"
GIGACHAT_API_URL = "https://gigachat.devices.sberbank.ru/api/v1/chat/completions"

token_cache = {"access_token": None, "expires_at": None}
TEMPLATE_FIELDS = [
    "Название",
    "Что хотим улучшить?",
    "Какие данные поступают агенту на выход?",
    "Как процесс выглядит сейчас? as-is",
    "Какой результат нужен от агента?",
    "Достижимый идеал(to-be)",
    "Масштаб процесса"
]
user_states = {}

# === Получение токена GigaChat ===
def get_gigachat_token():
    global token_cache
    if token_cache["access_token"] and token_cache["expires_at"] > datetime.utcnow():
        return token_cache["access_token"]

    headers = {
        'Content-Type': 'application/x-www-form-urlencoded',
        'Accept': 'application/json',
        'RqUID': str(uuid.uuid4()),
        'Authorization': f'Basic {GIGACHAT_AUTH_KEY}'
    }
    data = {'scope': GIGACHAT_SCOPE}
    response = requests.post(GIGACHAT_TOKEN_URL, headers=headers, data=data)
    if response.status_code == 200:
        result = response.json()
        token_cache["access_token"] = result['access_token']
        token_cache["expires_at"] = datetime.utcnow() + timedelta(seconds=result['expires_in'])
        return token_cache["access_token"]
    raise Exception("Не удалось получить токен GigaChat")

# === Проверка идеи через GigaChat ===
def check_idea_with_gigachat(user_input: str) -> str:
    try:
        wb = load_workbook("agents.xlsx")
        ws = wb.active
        all_agents_data = "\n".join([
            f"Название: {row[0]}, Команда: {row[1]}, Контакт: {row[2]}, Описание: {row[3]}"
            for row in ws.iter_rows(min_row=2, values_only=True) if row[0]
        ])
    except Exception as e:
        all_agents_data = "(не удалось загрузить данные об агентах)"

    prompt = (
        f"Вот список существующих AI-агентов:\n{all_agents_data}\n\n"
        f"Пользователь предлагает идею: {user_input}.\n"
        "Проверь, есть ли похожие идеи и оцени, насколько она уникальна. Ответь кратко и по делу."
    )

    access_token = get_gigachat_token()
    headers = {
        'Authorization': f'Bearer {access_token}',
        'Accept': 'application/json',
        'Content-Type': 'application/json'
    }
    data = {
        "model": "GigaChat-Pro",  # или GigaChat, в зависимости от доступных моделей
        "messages": [{"role": "user", "content": prompt}]
    }

    response = requests.post(GIGACHAT_API_URL, headers=headers, json=data)
    if response.status_code == 200:
        return response.json()['choices'][0]['message']['content']
    return f"Ошибка GigaChat API: {response.status_code}"

# === Генерация Word/Excel ===
def generate_files(data):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    word_path = f"agent_{timestamp}.docx"
    excel_path = f"agent_{timestamp}.xlsx"

    doc = Document()
    title = doc.add_heading("AI-агент — шаблон", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for key, value in data.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{key}:\n")
        run.bold = True
        run.font.size = Pt(14)
        run2 = p.add_run(f"{value}\n")
        run2.font.size = Pt(12)
        p.space_after = Pt(12)
    doc.save(word_path)

    wb = Workbook()
    ws = wb.active
    ws.title = "Агент"
    bold_font = Font(bold=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )
    alignment = Alignment(wrap_text=True, vertical="top")

    ws.append(["Поле", "Значение"])
    for cell in ws[1]:
        cell.font = bold_font
        cell.border = thin_border
        cell.alignment = alignment

    for key, value in data.items():
        ws.append([key, value])
        for cell in ws[ws.max_row]:
            cell.border = thin_border
            cell.alignment = alignment

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 60
    wb.save(excel_path)
    return word_path, excel_path

# === Основной роут ===
async def handler(request):
    data = await request.json()
    user_id = data['from']['userId']
    text = data['message']['text'].strip()

    if user_id not in user_states:
        user_states[user_id] = {}

    if text.lower() in ["/start", "привет"]:
        return web.json_response({"text": "👋 Привет! Опиши свою идею, и я проверю её уникальность."})

    if "step" not in user_states[user_id]:
        response = check_idea_with_gigachat(text)
        user_states[user_id]["last_idea"] = text
        user_states[user_id]["summary"] = response

        if any(word in response.lower() for word in ["уникальн"]):
            user_states[user_id]["step"] = 0
            user_states[user_id]["data"] = {TEMPLATE_FIELDS[0]: text}
            return web.json_response({"text": f"Идея выглядит уникальной! Давай заполним шаблон.\n1️⃣ {TEMPLATE_FIELDS[1]}:"})
        else:
            return web.json_response({"text": f"Похоже, такая идея уже есть.\n\n🤖 Ответ GigaChat:\n{response}"})
    else:
        state = user_states[user_id]
        step = state["step"] + 1
        state["data"][TEMPLATE_FIELDS[step]] = text
        if step + 1 < len(TEMPLATE_FIELDS):
            state["step"] = step
            return web.json_response({"text": f"{step+1}️⃣ {TEMPLATE_FIELDS[step+1]}:"})
        else:
            word_path, excel_path = generate_files(state["data"])
            del user_states[user_id]
            return web.json_response({"text": "✅ Файлы готовы. Спасибо!"})

# === Запуск сервера ===
logging.basicConfig(level=logging.INFO)
app = web.Application()
app.router.add_post("/gigabot", handler)

if __name__ == '__main__':
    web.run_app(app, port=8080)
