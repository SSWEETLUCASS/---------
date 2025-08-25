import re
import os
from datetime import datetime
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Border, Side, Alignment, PatternFill
from openpyxl.chart import BarChart, Reference
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from gigachat_wrapper import get_llm

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import io
import logging
from collections import defaultdict, deque
import json
from typing import Dict, List, Tuple, Optional, Any


# Настройка логирования
def setup_logging():
    """Настройка системы логирования"""
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
        handlers=[
            logging.FileHandler("gigachat.log", encoding="utf-8"),
            logging.StreamHandler()
        ]
    )

setup_logging()

# Память для диалогов с пользователями (user_id -> история последних 10 сообщений)
gigachat_memory = defaultdict(lambda: deque(maxlen=10))

class AgentDataProcessor:
    """Класс для работы с данными агентов"""
    
    @staticmethod
    def load_agents_data() -> List[Dict[str, str]]:
        """Загрузка данных об агентах из файла"""
        try:
            if not os.path.exists("agents.xlsx"):
                logging.warning("Файл agents.xlsx не найден")
                return []
                
            wb = load_workbook("agents.xlsx", data_only=True)
            ws = wb.active
            agents_data = []

            for row in ws.iter_rows(min_row=2, values_only=True):
                if not row or len(row) < 8 or not row[4]:
                    continue
                
                block, ssp, owner, contact, name, short_name, desc, typ = row[:8]
                agent_info = {
                    "block": str(block) if block else "",
                    "ssp": str(ssp) if ssp else "",
                    "owner": str(owner) if owner else "",
                    "contact": str(contact) if contact else "",
                    "name": str(name) if name else "",
                    "short_name": str(short_name) if short_name else "",
                    "description": str(desc) if desc else "",
                    "type": str(typ) if typ else ""
                }
                agents_data.append(agent_info)

            logging.info(f"Загружено {len(agents_data)} записей агентов")
            return agents_data
            
        except Exception as e:
            logging.error(f"Ошибка при загрузке agents.xlsx: {e}")
            return []

class MemoryManager:
    """Класс для управления памятью диалогов"""
    
    @staticmethod
    def add_to_memory(user_id: Optional[int], user_message: str, bot_response: str) -> None:
        if not user_id:
            return
        
        gigachat_memory[user_id].append({
            "timestamp": datetime.now().isoformat(timespec="seconds"),
            "user": user_message.strip()[:500],
            "bot": bot_response.strip()[:500]
        })

    @staticmethod
    def get_conversation_context(user_id: Optional[int]) -> str:
        if not user_id or user_id not in gigachat_memory:
            return ""
        
        history = list(gigachat_memory[user_id])
        if not history:
            return ""
        
        context_parts = []
        for i, exchange in enumerate(history, 1):
            context_parts.extend([
                f"Сообщение {i}:",
                f"Пользователь: {exchange['user']}",
                f"Бот: {exchange['bot']}",
                ""
            ])
        return "\n".join(context_parts)
    
    @staticmethod
    def clear_user_memory(user_id: int) -> bool:
        if user_id in gigachat_memory:
            gigachat_memory[user_id].clear()
            logging.info(f"Память пользователя {user_id} очищена")
            return True
        return False
    
    @staticmethod
    def get_memory_summary(user_id: int) -> str:
        if not user_id or user_id not in gigachat_memory:
            return "Память пуста"
        
        history = list(gigachat_memory[user_id])
        if not history:
            return "История диалога пуста"
        
        return f"В памяти {len(history)} обменов сообщениями. Последнее: {history[-1]['timestamp']}"

class CommandDetector:
    """Класс для детекции команд"""
    
    @staticmethod
    def detect_command(message: str) -> Optional[str]:
        message_lower = message.lower().strip()

        if any(word in message_lower for word in ['привет', 'hello', 'hi', 'начать', 'старт', '/start']):
            return 'start'
        if any(word in message_lower for word in ['помощь', 'help', 'справка', 'что умеешь', 'команды']):
            return 'help'
        if any(phrase in message_lower for phrase in ['список агентов', 'все агенты', 'покажи агентов', 'агенты список']):
            return 'ai_agent'
        if any(phrase in message_lower for phrase in ['найди агент', 'кто занимается', 'владелец агента', 'контакт']):
            return 'search_owners'
        if any(phrase in message_lower for phrase in ['придумай идею', 'посоветуй', 'идеи для', 'что делать с']):
            return 'consultation'

        # длинный текст → идея, короткий → уточнение (без команды)
        if (len(message) > 100 or any(phrase in message_lower for phrase in ['идея агента', 'создать агента'])):
            return 'idea'

        return None

class TextProcessor:
    """Класс для обработки текста"""
    
    @staticmethod
    def clean_response_text(text: str) -> str:
        """Улучшенная очистка текста ответа от служебных символов"""
        if not isinstance(text, str):
            text = str(text)
        
        # Убираем метаданные после content=
        if 'content=' in text:
            match = re.search(r"content=['\"]([^'\"]*)['\"]", text)
            if match:
                text = match.group(1)
            else:
                text = re.sub(r".*content=", "", text)
                text = re.sub(r"\s+additional_kwargs=.*$", "", text, flags=re.DOTALL)
        
        # Убираем дополнительные метаданные
        metadata_patterns = [
            r"\s*additional_kwargs=.*$",
            r"\s*response_metadata=.*$", 
            r"\s*id=.*$",
            r"\s*usage_metadata=.*$"
        ]
        for pattern in metadata_patterns:
            text = re.sub(pattern, "", text, flags=re.DOTALL)
        
        # Исправляем кодировку
        try:
            if isinstance(text, bytes):
                text = text.decode('utf-8')
            if 'Ð' in text or 'Ñ' in text:
                try:
                    text = text.encode('latin-1').decode('utf-8')
                except (UnicodeEncodeError, UnicodeDecodeError):
                    pass
        except Exception as e:
            logging.warning(f"Проблема с кодировкой: {e}")
        
        # Обработка escape-последовательностей
        escape_replacements = {
            '\\n': '\n',
            '\\t': '\t', 
            '\\"': '"',
            "\\'": "'"
        }
        for old, new in escape_replacements.items():
            text = text.replace(old, new)
        
        # Удаляем лишние слеши
        text = re.sub(r'\\(?![nrt"\'])', '', text)
        
        # Очищаем от служебных команд
        text = re.sub(r'^CMD:\w+\s*[•\-]*\s*', '', text)
        
        # Обработка форматирования
        text = re.sub(r'\s*--\s*', ' – ', text)
        text = re.sub(r'\s*##\s*', '\n\n', text)

        #  Преобразуем Markdown-жирный и подчёркнутый в *текст*
        text = re.sub(r'\*\*(.*?)\*\*', r'*\1*', text)  # **жирный** → *жирный*
        text = re.sub(r'__([^_]+)__', r'*\1*', text)    # __подчёркнутый__ → *подчёркнутый*

        # Финальная очистка
        text = text.strip()
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        return text


    @staticmethod
    def safe_str(value: Any) -> str:
        """Безопасное преобразование значений в строку"""
        if isinstance(value, dict):
            # Проверяем, не является ли это результатом интерактивного расчета
            if 'question' in value and 'key' in value:
                return f"Вопрос: {value['question']}"
            elif 'done' in value and 'result' in value:
                return value['result']
            return json.dumps(value, ensure_ascii=False, indent=2)
        if isinstance(value, list):
            return ", ".join(map(str, value))
        if value is None:
            return ""
        return str(value)

class GigaChatProcessor:
    """Класс для работы с GigaChat"""
    
    def __init__(self):
        self.agent_processor = AgentDataProcessor()
        self.memory_manager = MemoryManager()
        self.text_processor = TextProcessor()
        self.command_detector = CommandDetector()

    def check_idea_with_gigachat(self, user_input: str, user_data: Dict[str, Any], 
                                is_free_form: bool = False) -> Tuple[str, bool, Dict, bool]:
        """Проверка идеи с помощью GigaChat"""
        try:
            agents_data = self.agent_processor.load_agents_data()
            
            # Формируем данные об агентах
            if agents_data:
                joined_data = "\n\n".join([
                    f"Блок: {agent['block']}\n"
                    f"ССП: {agent['ssp']}\n"
                    f"Владелец: {agent['owner']}\n" 
                    f"Контакт: {agent['contact']}\n"
                    f"Название: {agent['name']}\n"
                    f"Краткое название: {agent['short_name']}\n"
                    f"Описание: {agent['description']}\n"
                    f"Тип: {agent['type']}"
                    for agent in agents_data
                ])
            else:
                joined_data = "(список инициатив пуст)"

            # Генерируем промпт в зависимости от типа ввода
            prompt = self._generate_idea_prompt(joined_data, user_data, is_free_form)
            
            # Отправляем запрос
            logging.info(f"[GigaChat Input] {prompt[:200]}...")
            raw_response = get_llm().invoke(prompt)
            logging.info(f"[GigaChat Output] {str(raw_response)[:200]}...")

            response_text = self.text_processor.clean_response_text(raw_response)

            # Анализируем ответ
            unclear_phrases = [
                "извините", "идея кажется не ясной", "идея не ясна",
                "идея глупая", "не очень хорошая идея", "давайте еще подумаем"
            ]
            
            unclear_idea = any(phrase in response_text.lower() for phrase in unclear_phrases)
            
            if unclear_idea:
                return response_text, False, {}, False

            # Сохраняем в память
            user_id = user_data.get("user_id")
            if user_id:
                self.memory_manager.add_to_memory(user_id, user_input, response_text)

            # Определяем уникальность
            is_unique = ("уникальна" in response_text.lower() and 
                        "не уникальна" not in response_text.lower())

            # Парсинг данных из свободной формы
            parsed_data = {}
            if is_free_form:
                parsed_data = self._parse_free_form_data(response_text)
                if is_unique and parsed_data:
                    try:
                        cost_calculator = CostCalculator()
                        cost = cost_calculator.calculate_work_cost_interactive(parsed_data)
                        # Исправляем обработку результата интерактивного расчета
                        if isinstance(cost, dict):
                            if 'result' in cost:
                                response_text += f"\n\n💰 {cost['result']}"
                        elif isinstance(cost, (int, float)) and cost > 0:
                            response_text += f"\n\n💰 Примерная стоимость работы: {cost:,.0f} ₽"
                    except Exception as e:
                        logging.error(f"Ошибка при расчете стоимости: {e}")

            suggest_processing = any(phrase in response_text.lower() for phrase in [
                "похоже на идею", "возможно, вы описали идею"
            ])

            return response_text, is_unique, parsed_data, suggest_processing

        except Exception as e:
            error_msg = f"⚠️ Ошибка при обращении к GigaChat: {e}"
            logging.error(error_msg)
            return error_msg, False, {}, False

    def check_general_message(self, msg: str, user_id: Optional[int] = None) -> Tuple[str, Optional[str]]:
        """Проверка общего сообщения - сначала локальная детекция команд, затем GigaChat"""
        try:
            # Сначала пробуем определить команду локально
            detected_command = self.command_detector.detect_command(msg)
            
            if detected_command:
                # Если команда определена локально, возвращаем её без обращения к GigaChat
                logging.info(f"[Local Command Detection] Detected: {detected_command}")
                return "", detected_command
            
            # Если команда не определена, обращаемся к GigaChat
            # Получаем контекст предыдущих сообщений
            conversation_context = self.memory_manager.get_conversation_context(user_id)
            
            prompt = self._generate_general_message_prompt(msg, conversation_context)
            
            logging.info(f"[GigaChat Input] {prompt[:200]}...")
            raw_response = get_llm().invoke(prompt)
            logging.info(f"[GigaChat Output] {str(raw_response)[:200]}...")

            response = self.text_processor.clean_response_text(raw_response)
            
            # Извлекаем команду
            cmd_match = re.search(r'CMD:(\w+)', response)
            detected_command = cmd_match.group(1) if cmd_match else None
            
            # Убираем команду из текста
            if cmd_match:
                response = re.sub(r'\s*CMD:\w+\s*', '', response).strip()
            
            # Сохраняем в память только если есть содержательный ответ
            if user_id and response:
                self.memory_manager.add_to_memory(user_id, msg, response)
            
            return response, detected_command
            
        except Exception as e:
            error_msg = f"⚠️ Ошибка при генерации ответа: {e}"
            logging.error(error_msg)
            return error_msg, None

    def find_agent_owners(self, query: str) -> str:
        """Поиск владельцев агентов по запросу"""
        try:
            agents_data = self.agent_processor.load_agents_data()
            
            if not agents_data:
                return "⚠️ Файл с агентами пуст или не найден."
            
            # Формируем информацию об агентах
            agents_info = "\n\n".join([
                f"Название: {agent['name']}\n"
                f"Описание: {agent['description']}\n"
                f"Тип: {agent['type']}\n"
                f"Блок: {agent['block']}\n"
                f"Владелец: {agent['owner']}\n"
                f"Контакт: {agent['contact']}"
                for agent in agents_data[:20]  # Ограничиваем количество для оптимизации
            ])
            
            prompt = f"""
            Запрос пользователя: "{query}"
            
            Доступные AI-агенты:
            {agents_info}
            
            Найди агентов, которые наиболее соответствуют запросу пользователя.
            Учитывай название, описание, тип и область применения.
            
            Для каждого подходящего агента выведи:
            - Название
            - Краткое описание  
            - Владелец и контакты
            
            Отвечай ТОЛЬКО на русском языке, используй смайлики для наглядности.
            """
            
            logging.info(f"[GigaChat Search Input] {query}")
            raw_response = get_llm().invoke(prompt)
            logging.info(f"[GigaChat Search Output] {str(raw_response)[:200]}...")
            
            response = self.text_processor.clean_response_text(raw_response)
            
            return response if response else "🤖 Не удалось найти подходящих агентов по вашему запросу."
            
        except Exception as e:
            error_msg = f"⚠️ Ошибка при поиске владельцев: {e}"
            logging.error(error_msg)
            return error_msg

    def _generate_idea_prompt(self, joined_data: str, user_data: Dict[str, Any], 
                             is_free_form: bool) -> str:
        """Генерация промпта для проверки идеи"""
        if is_free_form:
            return f"""
            Существующие инициативы:
            {joined_data}

            1. Проанализируй данный тебе текст пользователя и собери его по шаблону:
            - "Название"
            - "Что хотим улучшить?" 
            - "Какие данные поступают агенту на выход?"
            - "Как процесс выглядит сейчас? as-is"
            - "Какой результат нужен от агента?"
            - "Достижимый идеал(to-be)"
            - "Масштаб процесса"

            Если пользователь что-то не написал, укажи это и предложи уточнить.

            2. Сравни инициативу пользователя с существующими:
            - Если идея похожа на существующую — напиши "НЕ уникальна" и укажи название похожей инициативы и владельца.
            - Если идея новая — напиши "Уникальна" и предложи рекомендации по улучшению.
            - Если текст непонятный — напиши "Извините, не могу понять описание".

            3. Дай конструктивные рекомендации по развитию идеи.

            Отвечай ТОЛЬКО на русском языке, без дополнительной технической информации. Используй смайлики.

            Текст пользователя:
            \"\"\"{user_data.get('Описание в свободной форме', '')}\"\"\"
            """
        else:
            user_initiative = "\n".join([f"{key}: {value}" for key, value in user_data.items()])
            
            return f"""
            Инициатива пользователя:
            {user_initiative}

            Существующие инициативы:
            {joined_data}

            Задачи:
            1. Внимательно сравни инициативу пользователя с существующими инициативами.
            
            2. Определи уникальность:
            - Если идея похожа на существующую — напиши "НЕ уникальна" и укажи название похожей инициативы и владельца.
            - Если идея новая — напиши "Уникальна" и предложи рекомендации по улучшению.
            
            3. Дай детальную оценку инициативы и советы по её развитию.

            4. Если идея кажется неясной или требует доработки, пиши: "Извините, но давайте еще подумаем".

            Отвечай ТОЛЬКО на русском языке, без дополнительной технической информации. Используй смайлики.
            """

    def _generate_general_message_prompt(self, msg: str, conversation_context: str) -> str:
        """Генерация промпта для общего сообщения"""
        context_section = ""
        if conversation_context:
            context_section = f"""
            История нашего диалога:
            {conversation_context}

            Текущее сообщение пользователя:
            """

        return f"""{context_section}
        Пользователь написал:
        \"\"\"{msg}\"\"\"
Контекст: Ты - помощник по разработке AI-агентов. Учитывай предыдущие сообщения пользователя для более конструктивного диалога.

Твоя задача — понять смысл сообщения и определить подходящую команду для бота.

Правила выбора команды:
1. Приветствие или начало общения → CMD:start
2. Если пользователь только упомянул, что у него есть идея (без деталей) → начинай пошаговое уточнение по шаблону (без CMD).
3. Первичное описание идеи с достаточным количеством информации (≥70% всех полей шаблона заполнены) → CMD:idea
4. Если идея неполная (<100%), пошагово спрашивай недостающие поля из шаблона:
   - "Название"
   - "Что хотим улучшить?"
   - "Какие данные поступают агенту на выход?"
   - "Как процесс выглядит сейчас? as-is"
   - "Какой результат нужен от агента?"
   - "Достижимый идеал (to-be)"
   - "Масштаб процесса"
5. Уточнение по уже существующей идее → остаёмся в процессе (без CMD).
6. Жалобы на проблемы с ботом → CMD:help
7. Запрос списка агентов → CMD:ai_agent
8. Консультация (ссылки, рекомендации) → CMD:consultation
9. Вопросы про владельцев → CMD:search_owners
10. Иначе — полезный ответ без команды

Особенности ответа:
- Если пользователь написал только «есть идея» или «хочу поделиться мыслью» → спроси: «Какое название у вашей идеи?».
- Если заполнение ≥70%, сразу переходи к CMD:idea.
- Если <70%, уточняй по каждому пункту из шаблона, один за другим, пока не будет заполнено всё.
- Когда заполнен весь шаблон → выводи итог в виде собранной идеи и указывай CMD:idea.
- Команду вставляй в ответ только один раз (если это новый процесс).
- Поддерживай непрерывность диалога и дружелюбный тон со смайликами.
- Не более 4000 символов.

Формат: [Текст ответа] [CMD:команда] (если применимо)
Отвечай ТОЛЬКО на русском языке.


        """

    def _parse_free_form_data(self, response_text: str) -> Dict[str, str]:
        """Парсинг данных из свободной формы ответа"""
        parsed_data = {}
        fields = [
            "Название", "Что хотим улучшить?", "Какие данные поступают агенту на выход?",
            "Как процесс выглядит сейчас? as-is", "Какой результат нужен от агента?",
            "Достижимый идеал(to-be)", "Масштаб процесса"
        ]
        
        for field in fields:
            patterns = [
                rf"{re.escape(field)}[:\-–]\s*(.+?)(?=\n\w+[:\-–]|$)",
                rf"{re.escape(field.lower())}[:\-–]\s*(.+?)(?=\n\w+[:\-–]|$)",
            ]
            for pattern in patterns:
                match = re.search(pattern, response_text, re.IGNORECASE | re.DOTALL)
                if match:
                    parsed_data[field] = match.group(1).strip()
                    break
        
        return parsed_data

class CostCalculator:
    """Класс для расчета стоимости проектов"""
    
    def __init__(self):
        self.hourly_rate = 3500
        self.base_hours = 40
        self.scale_map = {
            "малый": 1, "мал": 1, "небольшой": 1,
            "средний": 1.8, "средн": 1.8,  
            "большой": 2.5, "больш": 2.5,
            "крупный": 3.2, "крупн": 3.2,
            "глобальный": 4, "глобальн": 4, "масштабный": 4
        }

    def calculate_work_cost(self, parsed_data: Dict[str, Any], is_unique: bool = True) -> str:
        """Расчет примерной стоимости работы по инициативе"""
        try:
            # Определяем коэффициент масштаба
            scale_value = str(parsed_data.get("Масштаб процесса", "")).strip().lower()
            
            if scale_value.replace('.', '').replace(',', '').isdigit():
                hours_coefficient = min(max(float(scale_value.replace(',', '.')), 0.5), 5.0)
            else:
                hours_coefficient = 1.0
                for key, value in self.scale_map.items():
                    if key in scale_value:
                        hours_coefficient = value
                        break

            # Анализ сложности
            description_text = " ".join([
                str(parsed_data.get("Описание", "")),
                str(parsed_data.get("Как процесс выглядит сейчас? as-is", "")),
                str(parsed_data.get("Какой результат нужен от агента?", ""))
            ]).lower()

            complexity_bonus = 0
            
            # Ключевые слова сложности
            complex_keywords = [
                "интеграция", "апи", "api", "машинное обучение", "ml", "ai", 
                "нейронн", "алгоритм", "распознавание", "nlp", "компьютерное зрение",
                "большие данные", "реальное время", "безопасность", "криптография"
            ]
            
            simple_keywords = ["простой", "базовый", "стандартн", "типовой", "шаблон"]
            
            for keyword in complex_keywords:
                if keyword in description_text:
                    complexity_bonus += 0.3
                    
            for keyword in simple_keywords:
                if keyword in description_text:
                    complexity_bonus -= 0.2

            complexity_bonus = max(-0.5, min(complexity_bonus, 1.5))
            uniqueness_coefficient = 1.0 if is_unique else 0.7

            # Расчет часов и стоимости
            total_hours = max(20, self.base_hours * hours_coefficient * 
                            (1 + complexity_bonus) * uniqueness_coefficient)
            
            analysis_hours = total_hours * 0.15
            development_hours = total_hours * 0.60
            testing_hours = total_hours * 0.15
            deployment_hours = total_hours * 0.10
            
            total_cost = total_hours * self.hourly_rate

            return f"""
📊 **Детальный расчет стоимости разработки:**

🔢 **Трудозатраты:**
• Анализ и проектирование: {analysis_hours:.0f} ч.
• Разработка и программирование: {development_hours:.0f} ч.
• Тестирование и отладка: {testing_hours:.0f} ч.
• Внедрение и документация: {deployment_hours:.0f} ч.
**Всего часов: {total_hours:.0f} ч.**

💰 **Финансовые расчеты:**
• Ставка разработчика: {self.hourly_rate:,} ₽/час
• Коэффициент масштаба: {hours_coefficient}x
• Коэффициент сложности: {(1 + complexity_bonus):.2f}x
• Коэффициент уникальности: {uniqueness_coefficient}x
• Уникальность идеи: {'Да' if is_unique else 'Нет (есть аналоги)'}

💸 **ИТОГОВАЯ СТОИМОСТЬ: {total_cost:,.0f} ₽**
💼 **({total_hours:.0f} чел./час)**

📈 **Диапазон стоимости:** {total_cost*0.8:,.0f} - {total_cost*1.3:,.0f} ₽

📝 **Примечание:** Стоимость может изменяться в зависимости от детальных требований.
            """
        except Exception as e:
            logging.error(f"Ошибка при расчете стоимости: {e}")
            return f"⚠️ Ошибка при расчете стоимости: {e}"

    def calculate_work_cost_interactive(self, answers: Dict[str, Any], return_next: bool = False) -> Any:
        """Интерактивный расчет стоимости"""
        questions = [
            ("Название", "Как называется ваша инициатива?"),
            ("Что хотим улучшить?", "Что именно вы хотите улучшить с помощью агента?"),
            ("Какие данные поступают агенту на выход?", "Какие данные агент будет выдавать на выходе?"),
            ("Масштаб процесса", "Каков масштаб процесса (малый, средний, большой)?"),
        ]

        # Поиск следующего вопроса
        for key, question in questions:
            if key not in answers or answers[key] is None:
                if return_next:
                    return {"question": question, "key": key}
                answers[key] = None

        # Если все ответы есть — считаем стоимость
        try:
            cost_description = self.calculate_work_cost(answers)
            # Извлекаем только числовое значение
            cost_match = re.search(r'ИТОГОВАЯ СТОИМОСТЬ:\s*([\d,]+)', cost_description)
            if cost_match:
                cost_value = int(cost_match.group(1).replace(',', ''))
                if return_next:
                    return {"done": True, "result": f"Примерная стоимость: {cost_value:,.0f} ₽"}
                return cost_value
        except Exception as e:
            logging.error(f"Ошибка в интерактивном расчете: {e}")
            
        if return_next:
            return {"done": True, "result": "Не удалось рассчитать стоимость"}
        return 0

class FileGenerator:
    """Класс для генерации файлов"""
    
    @staticmethod
    def generate_files(data: Dict[str, Any], cost_info: str = "") -> Tuple[str, str]:
        """Генерация Word и Excel файлов с данными инициативы"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        word_path = f"initiative_{timestamp}.docx"
        excel_path = f"initiative_{timestamp}.xlsx"

        try:
            # Создание Word документа
            doc = Document()
            
            # Заголовок
            title = doc.add_heading("Описание AI-инициативы", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Дата создания
            date_para = doc.add_paragraph(f"Дата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.runs[0]
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # Пустая строка
            
            # Основной контент
            for key, value in data.items():
                # Заголовок поля
                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(f"📋 {key}")
                heading_run.bold = True
                heading_run.font.size = Pt(14)
                heading_run.font.color.rgb = RGBColor(0, 70, 132)
                
                # Разделительная линия
                line_para = doc.add_paragraph("─" * 50)
                line_run = line_para.runs[0]
                line_run.font.color.rgb = RGBColor(200, 200, 200)
                
                # Содержимое поля
                content_para = doc.add_paragraph(str(value))
                content_run = content_para.runs[0]
                content_run.font.size = Pt(12)
                
                doc.add_paragraph()  # Пустая строка между разделами
            
            # Добавляем информацию о стоимости
            if cost_info:
                cost_heading = doc.add_heading("💰 Расчет стоимости", level=1)
                cost_para = doc.add_paragraph(cost_info)
                cost_run = cost_para.runs[0]
                cost_run.font.size = Pt(11)
            
            # Футер
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run("Создано с помощью Агентолога 🤖")
            footer_run.font.size = Pt(10)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.save(word_path)

            # Создание Excel файла
            wb = Workbook()
            ws = wb.active
            ws.title = "Инициатива"

            # Стили
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            thin_border = Border(
                left=Side(style="thin"), right=Side(style="thin"),
                top=Side(style="thin"), bottom=Side(style="thin")
            )
            wrap_alignment = Alignment(wrap_text=True, vertical="top")
            
            # Заголовки
            ws.append(["Поле", "Значение"])
            for cell in ws[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border
                cell.alignment = Alignment(horizontal="center", vertical="center")
            
            # Данные
            for key, value in data.items():
                ws.append([key, str(value)])
                for cell in ws[ws.max_row]:
                    cell.border = thin_border
                    cell.alignment = wrap_alignment
            
            # Добавляем информацию о стоимости в Excel
            if cost_info:
                ws.append(["", ""])  # Пустая строка
                ws.append(["Расчет стоимости", cost_info])
                for cell in ws[ws.max_row]:
                    cell.border = thin_border
                    cell.alignment = wrap_alignment
            
            # Форматирование колонок
            ws.column_dimensions["A"].width = 35
            ws.column_dimensions["B"].width = 70
            
            # Информационная строка
            ws.append(["", ""])
            info_row = ws.max_row + 1
            ws[f"A{info_row}"] = "Создано"
            ws[f"B{info_row}"] = datetime.now().strftime('%d.%m.%Y %H:%M')
            
            for cell in [ws[f"A{info_row}"], ws[f"B{info_row}"]]:
                cell.font = Font(italic=True, color="808080")
                cell.border = thin_border
            
            wb.save(excel_path)
            
            logging.info(f"Файлы созданы: {word_path}, {excel_path}")
            return word_path, excel_path
            
        except Exception as e:
            logging.error(f"Ошибка при создании файлов: {e}")
            raise

    @staticmethod
    def generate_agents_summary_file(agents_file_path: str) -> Optional[str]:
        """Генерация аналитического файла с агентами"""
        try:
            agent_processor = AgentDataProcessor()
            agents_data = agent_processor.load_agents_data()
            
            if not agents_data:
                return None
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            summary_file = f"agents_summary_{timestamp}.xlsx"
            
            wb = Workbook()
            
            # Лист 1: Исходные данные
            FileGenerator._create_agents_list_sheet(wb, agents_data)
            
            # Лист 2: Аналитика
            FileGenerator._create_analytics_sheet(wb, agents_data)
            
            # Лист 3: Контакты
            FileGenerator._create_contacts_sheet(wb, agents_data)
            
            wb.save(summary_file)
            logging.info(f"Аналитический файл создан: {summary_file}")
            return summary_file
            
        except Exception as e:
            logging.error(f"Ошибка при создании аналитического файла: {e}")
            return None

    @staticmethod
    def _create_agents_list_sheet(wb: Workbook, agents_data: List[Dict[str, str]]) -> None:
        """Создание листа со списком агентов"""
        ws1 = wb.active
        ws1.title = "Список агентов"
        
        # Заголовки
        headers = ["Блок", "ССП", "Владелец", "Контакт", "Название", "Краткое название", "Описание", "Тип"]
        ws1.append(headers)
        
        # Стили для заголовков
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin")
        )
        
        for cell in ws1[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Добавляем данные
        for agent in agents_data:
            ws1.append([
                agent['block'], agent['ssp'], agent['owner'], agent['contact'],
                agent['name'], agent['short_name'], agent['description'], agent['type']
            ])
        
        # Форматирование данных
        for row in ws1.iter_rows(min_row=2, max_row=ws1.max_row):
            for cell in row:
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        
        # Автоширина колонок
        for column in ws1.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    cell_length = len(str(cell.value)) if cell.value else 0
                    if cell_length > max_length:
                        max_length = cell_length
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws1.column_dimensions[column_letter].width = adjusted_width

    @staticmethod
    def _create_analytics_sheet(wb: Workbook, agents_data: List[Dict[str, str]]) -> None:
        """Создание аналитического листа"""
        ws2 = wb.create_sheet("Аналитика")
        
        # Статистика
        type_stats = defaultdict(int)
        block_stats = defaultdict(int)
        
        for agent in agents_data:
            agent_type = agent['type'] or "Не указан"
            agent_block = agent['block'] or "Не указан"
            
            type_stats[agent_type] += 1
            block_stats[agent_block] += 1
        
        # Заголовки аналитики
        ws2['A1'] = "Аналитический отчет по AI-агентам"
        ws2['A1'].font = Font(size=16, bold=True)
        ws2['A1'].alignment = Alignment(horizontal="center")
        ws2.merge_cells('A1:D1')
        
        # Общая статистика
        ws2['A3'] = "Общая статистика:"
        ws2['A3'].font = Font(bold=True, size=12)
        ws2['A4'] = f"Всего агентов: {len(agents_data)}"
        ws2['A5'] = f"Уникальных типов: {len(type_stats)}"
        ws2['A6'] = f"Уникальных блоков: {len(block_stats)}"
        
        # Статистика по типам
        ws2['A8'] = "Распределение по типам:"
        ws2['A8'].font = Font(bold=True, size=12)
        row = 9
        for agent_type, count in sorted(type_stats.items(), key=lambda x: x[1], reverse=True):
            ws2[f'A{row}'] = agent_type
            ws2[f'B{row}'] = count
            row += 1
        
        # Статистика по блокам
        ws2['D8'] = "Распределение по блокам:"
        ws2['D8'].font = Font(bold=True, size=12)
        row = 9
        for block, count in sorted(block_stats.items(), key=lambda x: x[1], reverse=True):
            ws2[f'D{row}'] = block
            ws2[f'E{row}'] = count
            row += 1

    @staticmethod
    def _create_contacts_sheet(wb: Workbook, agents_data: List[Dict[str, str]]) -> None:
        """Создание листа с контактами"""
        ws3 = wb.create_sheet("Контакты владельцев")
        ws3.append(["Владелец", "Контакт", "Количество агентов", "Названия агентов"])
        
        # Группируем по владельцам
        owner_agents = defaultdict(list)
        owner_contacts = {}
        
        for agent in agents_data:
            owner = agent['owner'] or "Не указан"
            owner_agents[owner].append(agent['name'])
            if not owner_contacts.get(owner):
                owner_contacts[owner] = agent['contact']
        
        for owner, agent_names in owner_agents.items():
            contact = owner_contacts.get(owner, "")
            ws3.append([owner, contact, len(agent_names), "; ".join(agent_names)])
        
        # Форматирование листа контактов
        for cell in ws3[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="70AD47", end_color="70AD47", fill_type="solid")

class DiagramGenerator:
    """Класс для генерации диаграмм"""
    
    @staticmethod
    def generate_idea_evaluation_diagram(idea_data: Dict[str, Any], is_unique: bool, 
                                       parsed_data: Optional[Dict] = None) -> Optional[str]:
        """Генерация паутинчатой диаграммы оценки идеи"""
        try:
            # Настройка шрифтов для кириллицы
            plt.rcParams['font.family'] = ['DejaVu Sans', 'Arial', 'sans-serif']
            plt.rcParams['axes.unicode_minus'] = False

            # Получаем оценки от GigaChat
            scores = DiagramGenerator._get_idea_scores(idea_data, parsed_data)
            
            # Построение диаграммы
            categories = list(scores.keys())
            values = list(scores.values())
            values += values[:1]  # замкнуть график

            angles = [n / float(len(categories)) * 2 * np.pi for n in range(len(categories))]
            angles += angles[:1]

            fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(polar=True))
            
            # Заголовок
            title_text = (parsed_data or idea_data).get("Название", "Новая идея")
            fig.suptitle(f'📊 Оценка AI-инициативы: {title_text}', 
                        fontsize=16, fontweight='bold', y=0.98)

            # Настройка осей
            ax.set_theta_offset(np.pi / 2)
            ax.set_theta_direction(-1)

            # Основной график
            ax.plot(angles, values, 'o-', linewidth=3, label='Оценка', color='#2E86C1', markersize=8)
            ax.fill(angles, values, alpha=0.25, color='#2E86C1')

            # Настройка сетки
            ax.set_xticks(angles[:-1])
            ax.set_xticklabels(categories, fontsize=11, fontweight='bold')
            ax.set_ylim(0, 10)
            ax.set_yticks([2, 4, 6, 8, 10])
            ax.set_yticklabels(['2', '4', '6', '8', '10'], fontsize=9)
            ax.grid(True, alpha=0.7)

            # Добавляем значения на точки
            for angle, value, category in zip(angles[:-1], values[:-1], categories):
                ax.annotate(f'{value}', xy=(angle, value), xytext=(5, 5), 
                           textcoords='offset points', fontsize=10, fontweight='bold',
                           bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.8))

            # Статистика
            avg_score = sum(scores.values()) / len(scores)
            status, status_color = DiagramGenerator._get_status_info(avg_score)
            uniqueness_text = "✅ Уникальная" if is_unique else "⚠️ Есть аналоги"
            
            info_text = f"Средняя оценка: {avg_score:.1f}/10  •  {status}  •  {uniqueness_text}"
            fig.text(0.5, 0.08, info_text, ha='center', fontsize=12, fontweight='bold',
                    bbox=dict(boxstyle="round,pad=0.7", facecolor=status_color, alpha=0.2))

            # Сохранение
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"idea_radar_{timestamp}.png"
            plt.savefig(filename, dpi=200, bbox_inches='tight', facecolor='white', 
                       edgecolor='none', pad_inches=0.3)
            plt.close()

            logging.info(f"Диаграмма создана: {filename}")
            return filename

        except Exception as e:
            logging.error(f"Ошибка при создании диаграммы: {e}")
            plt.close()  # Закрываем figure в случае ошибки
            return None

    @staticmethod 
    def _get_idea_scores(idea_data: Dict[str, Any], parsed_data: Optional[Dict] = None) -> Dict[str, int]:
        """Получение оценок идеи от GigaChat"""
        try:
            text_processor = TextProcessor()
            analysis_text = "\n".join([
                f"{k}: {text_processor.safe_str(v)}" 
                for k, v in (parsed_data or idea_data).items()
            ])

            evaluation_prompt = f"""
            Проанализируй следующую идею AI-агента и оцени её по 6 критериям от 1 до 10:

            Идея:
            {analysis_text}

            Критерии оценки:
            1. Актуальность (насколько проблема востребована сейчас)
            2. Сложность реализации (10 - очень сложно, 1 - очень просто)
            3. ROI потенциал (возврат инвестиций, экономический эффект)
            4. Инновационность (насколько идея новаторская)
            5. Масштабируемость (возможность расширения и тиражирования)
            6. Техническая осуществимость (реально ли это сделать с текущими технологиями)

            Отвечай СТРОГО в формате:
            Актуальность: X
            Сложность: X
            ROI: X
            Инновационность: X
            Масштабируемость: X
            Осуществимость: X
            
            Где X - число от 1 до 10.
            """
            
            raw_response = get_llm().invoke(evaluation_prompt)
            evaluation_text = text_processor.clean_response_text(raw_response)

            # Парсим оценки
            default_scores = {
                'Актуальность': 7, 'Сложность': 6, 'ROI': 6,
                'Инновационность': 5, 'Масштабируемость': 6, 'Осуществимость': 7
            }
            
            scores = {}
            for key, default_value in default_scores.items():
                match = re.search(rf"{key}[:\-–]\s*(\d+)", evaluation_text, re.IGNORECASE)
                if match:
                    score = min(max(int(match.group(1)), 1), 10)
                    scores[key] = score
                else:
                    scores[key] = default_value
                    
            return scores
            
        except Exception as e:
            logging.error(f"Ошибка при получении оценок: {e}")
            return {
                'Актуальность': 7, 'Сложность': 6, 'ROI': 6,
                'Инновационность': 5, 'Масштабируемость': 6, 'Осуществимость': 7
            }

    @staticmethod
    def _get_status_info(avg_score: float) -> Tuple[str, str]:
        """Определение статуса проекта по средней оценке"""
        if avg_score >= 7:
            return "🟢 РЕКОМЕНДУЕТСЯ", '#27AE60'
        elif avg_score >= 5:
            return "🟡 ДОРАБОТАТЬ", '#F39C12'
        else:
            return "🔴 РИСКИ", '#E74C3C'

class CostCalculationManager:
    """Класс для управления процессом расчета стоимости"""
    
    def __init__(self):
        self.gigachat_processor = GigaChatProcessor()
        self.text_processor = TextProcessor()

    def generate_cost_questions(self, parsed_data: Dict[str, Any]) -> Tuple[str, Optional[Dict]]:
        """Генерирует уточняющие вопросы для точного расчета стоимости"""
        try:
            initiative_context = "\n".join([f"{key}: {value}" for key, value in parsed_data.items()])
            
            prompt = f"""
            Проанализируй следующую AI-инициативу и сформируй 7-8 конкретных вопросов для точного расчета стоимости:

            ИНИЦИАТИВА:
            {initiative_context}

            Сформируй вопросы по аспектам:
            1. Команда разработки (сколько человек, роли)
            2. Временные рамки (дедлайны, этапы)
            3. Техническая сложность (интеграции, технологии)
            4. Объем данных и нагрузка
            5. Требования к качеству и безопасности
            6. Инфраструктура и развертывание
            7. Сопровождение и поддержка
            8. Дополнительные требования

            Каждый вопрос должен быть конкретным с вариантами ответов.

            Формат ответа:
            ВОПРОС 1: [текст вопроса]
            Варианты: [варианты ответов]

            Отвечай ТОЛЬКО на русском языке, добавь эмодзи.
            """
            
            logging.info(f"[GigaChat Questions] Generating cost questions")
            raw_response = get_llm().invoke(prompt)
            
            questions_text = self.text_processor.clean_response_text(raw_response)
            questions_dict = self._parse_questions_from_text(questions_text)
            
            response_text = f"""
🎯 **Для точного расчета стоимости мне нужно уточнить несколько деталей:**

{questions_text}

📝 **Как отвечать:**
Просто напишите номер вопроса и ваш ответ, например:
"1. 3 человека" или "2. 2 месяца"

Можно отвечать по несколько вопросов сразу или по одному.
            """
            
            return response_text, questions_dict
            
        except Exception as e:
            logging.error(f"Ошибка при генерации вопросов: {e}")
            return f"⚠️ Ошибка при генерации вопросов: {e}", None

    def process_cost_answers(self, questions: Dict, user_input: str) -> Tuple[Dict, bool, str]:
        """Обработка ответов пользователя на вопросы о стоимости"""
        try:
            # Парсим ответы
            answer_pattern = r'(\d+)\.?\s*(.+?)(?=\n\d+\.|\n|$)'
            matches = re.findall(answer_pattern, user_input, re.MULTILINE)
            
            answered_count = 0
            total_questions = len(questions)
            
            for match in matches:
                question_num = match[0]
                answer = match[1].strip()
                
                if question_num in questions:
                    questions[question_num]['answered'] = True
                    questions[question_num]['answer'] = answer
                    answered_count += 1
            
            # Проверяем статус
            all_answered = all(q['answered'] for q in questions.values())
            
            if answered_count == 0:
                status_msg = "❌ Не удалось распознать ответы. Используйте формат: '1. ваш ответ'"
            elif all_answered:
                status_msg = f"✅ Все {total_questions} вопросов отвечены! Делаю расчет..."
            else:
                answered_nums = [k for k, v in questions.items() if v['answered']]
                unanswered_nums = [k for k, v in questions.items() if not v['answered']]
                status_msg = (f"📝 Получил ответы на вопросы: {', '.join(answered_nums)}\n"
                            f"🔄 Остались вопросы: {', '.join(unanswered_nums)}\n\n"
                            f"Можете продолжить отвечать или написать 'рассчитать' для расчета.")
            
            return questions, all_answered, status_msg
            
        except Exception as e:
            logging.error(f"Ошибка при обработке ответов: {e}")
            return questions, False, f"⚠️ Ошибка при обработке ответов: {e}"

    def calculate_final_cost(self, parsed_data: Dict[str, Any], answers: Dict[str, str], 
                           user_id: Optional[int] = None) -> Tuple[str, Optional[Dict]]:
        """Финальный расчет стоимости на основе ответов"""
        try:
            initiative_context = "\n".join([f"{key}: {value}" for key, value in parsed_data.items()])
            answers_context = "\n".join([f"Вопрос {k}: {v}" for k, v in answers.items()])
            
            prompt = f"""
            Сделай детальный расчет стоимости разработки AI-агента:

            ИНИЦИАТИВА:
            {initiative_context}

            ОТВЕТЫ НА УТОЧНЯЮЩИЕ ВОПРОСЫ:
            {answers_context}

            ЗАДАЧА: Рассчитай реалистичную стоимость с учетом факторов:

            1. **Состав команды и роли:**
            - Аналитик/Product Owner
            - Backend разработчик
            - Frontend разработчик (если нужен UI)
            - Data Scientist/ML Engineer (если нужно ML)
            - DevOps инженер
            - QA инженер
            - Проект-менеджер

            2. **Трудозатраты по этапам:**
            - Анализ и проектирование
            - Разработка MVP
            - Тестирование и отладка
            - Интеграция и развертывание
            - Документация и обучение

            3. **Дополнительные расходы:**
            - Инфраструктура
            - Лицензии на ПО
            - Сторонние API/сервисы
            - Непредвиденные расходы (10-20%)

            **Используй ставки (₽/час):**
            Junior: 2000-3000, Middle: 3500-5000, Senior: 5500-7500, Lead: 7000-10000

            **ФОРМАТ ОТВЕТА:**
            👥 **СОСТАВ КОМАНДЫ:**
            [Роль] - [количество] - [уровень] - [ставка ₽/час]

            ⏱️ **ВРЕМЕННЫЕ ЗАТРАТЫ:**
            [Этап] - [часы] - [стоимость ₽]

            💰 **ИТОГОВАЯ СМЕТА:**
            Разработка: [сумма] ₽
            Инфраструктура: [сумма] ₽
            Дополнительные расходы: [сумма] ₽
            **ОБЩАЯ СТОИМОСТЬ: [итоговая сумма] ₽**

            📊 **ВРЕМЕННЫЕ РАМКИ:**
            Общее время: [X] месяцев
            Человеко-часов: [X] часов

            Будь конкретным и реалистичным!
            """
            
            logging.info(f"[GigaChat Final Cost] Calculating...")
            raw_response = get_llm().invoke(prompt)
            
            cost_calculation = self.text_processor.clean_response_text(raw_response)
            
            # Сохраняем в память
            if user_id:
                MemoryManager.add_to_memory(user_id, 
                    f"Расчет стоимости для: {parsed_data.get('Название', 'инициативы')}", 
                    cost_calculation)
            
            return cost_calculation, None
            
        except Exception as e:
            logging.error(f"Ошибка при финальном расчете: {e}")
            return f"⚠️ Ошибка при расчете стоимости: {e}", None

    def handle_cost_calculation_flow(self, user_input: str, user_data: Dict[str, Any], 
                                   user_id: Optional[int] = None) -> Tuple[str, Dict]:
        """Обработка флоу интерактивного расчета стоимости"""
        cost_state = user_data.get('cost_calculation_state', {})
        
        # Первый запрос на расчет
        if not cost_state:
            response, questions = self.generate_cost_questions(user_data)
            cost_state = {
                'stage': 'questions',
                'questions': questions,
                'start_time': datetime.now().isoformat()
            }
            return response, cost_state
        
        # Обработка ответов на вопросы
        if cost_state.get('stage') == 'questions':
            questions = cost_state.get('questions', {})
            
            # Проверяем принудительный расчет
            if any(word in user_input.lower() for word in ['рассчитать', 'посчитать', 'расчет']):
                answers = {k: v['answer'] for k, v in questions.items() if v.get('answered')}
                if answers:
                    final_cost, _ = self.calculate_final_cost(user_data, answers, user_id)
                    cost_state = {'stage': 'completed'}
                    return final_cost, cost_state
                else:
                    return "❌ Нет ответов для расчета. Ответьте хотя бы на несколько вопросов.", cost_state
            
            # Обрабатываем ответы
            updated_questions, all_answered, status_msg = self.process_cost_answers(questions, user_input)
            cost_state['questions'] = updated_questions
            
            if all_answered:
                answers = {k: v['answer'] for k, v in updated_questions.items()}
                final_cost, _ = self.calculate_final_cost(user_data, answers, user_id)
                cost_state = {'stage': 'completed'}
                return final_cost, cost_state
            else:
                return status_msg, cost_state
        
        # Расчет завершен
        if cost_state.get('stage') == 'completed':
            return "✅ Расчет стоимости завершен. Для нового расчета создайте новую инициативу.", cost_state
        
        return "⚠️ Неизвестное состояние расчета.", cost_state

    def _parse_questions_from_text(self, text: str) -> Dict[str, Dict]:
        """Парсинг вопросов из текста в структурированный словарь"""
        questions = {}
        
        question_pattern = r'ВОПРОС\s*(\d+):\s*(.+?)(?=\n|Варианты:|$)'
        variants_pattern = r'Варианты:\s*(.+?)(?=\n\s*ВОПРОС|\n\s*$|$)'
        
        question_matches = re.findall(question_pattern, text, re.DOTALL | re.IGNORECASE)
        
        for match in question_matches:
            question_num = match[0]
            question_text = match[1].strip()
            
            # Ищем варианты для этого вопроса
            question_block = re.search(
                rf'ВОПРОС\s*{question_num}:.*?(?=ВОПРОС\s*\d+:|$)', 
                text, re.DOTALL | re.IGNORECASE
            )
            
            variants = []
            if question_block:
                variants_match = re.search(variants_pattern, question_block.group(), re.DOTALL | re.IGNORECASE)
                if variants_match:
                    variants_text = variants_match.group(1).strip()
                    variants = [v.strip() for v in variants_text.split(',') if v.strip()]
            
            questions[question_num] = {
                'question': question_text,
                'variants': variants,
                'answered': False,
                'answer': None
            }
        
        return questions

class AIAgentBot:
    def __init__(self):
        self.gigachat = GigaChatProcessor()
        self.cost_manager = CostCalculationManager()
        self.file_generator = FileGenerator()
        self.diagram_generator = DiagramGenerator()
        self.memory_manager = MemoryManager()

    def process_user_message(self, message: str, user_id: Optional[int] = None, context: Optional[Dict] = None) -> Dict[str, Any]:
        try:
            response, command = self.gigachat.check_general_message(message, user_id)

            result = {
                'response': response or "",
                'command': command,
                'files': [],
                'diagram': None,
                'success': True
            }

            if command:
                cmd_result = self._handle_command(command, message, user_id, context)
                if 'response' in cmd_result:
                    result['response'] = (result['response'] + "\n" + cmd_result['response']).strip()
                result.update({k: v for k, v in cmd_result.items() if k != 'response'})
            return result
        except Exception as e:
            logging.error(f"Ошибка в обработке сообщения: {e}")
            return {'response': f"⚠️ Ошибка: {e}", 'command': None, 'files': [], 'diagram': None, 'success': False}

    def _handle_command(self, command: str, message: str, user_id: Optional[int], context: Optional[Dict]) -> Dict[str, Any]:
        result = {}
        if command == 'start':
            result['response'] = self._get_start_message()
        elif command == 'help':
            result['response'] = self._get_help_message()
        elif command == 'ai_agent':
            summary_file = self.get_agents_summary()
            if summary_file:
                result['files'] = [summary_file]
                result['response'] = f"📊 Создан аналитический файл: {summary_file}"
        elif command == 'search_owners':
            result['response'] = self.gigachat.find_agent_owners(message)
        elif command == 'consultation':
            result['response'] = self.generate_idea_suggestions(message)
        elif command == 'idea':
            result['response'] = "💡 Для обработки идеи используйте метод process_idea"
        return result


    def process_idea(self, user_data: Dict[str, Any], is_free_form: bool = False, 
                    user_id: Optional[int] = None) -> Dict[str, Any]:
        """Обработка идеи пользователя"""
        try:
            user_input = user_data.get('Описание в свободной форме', '') if is_free_form else str(user_data)
            
            # Анализируем идею
            response, is_unique, parsed_data, suggest_processing = self.gigachat.check_idea_with_gigachat(
                user_input, user_data, is_free_form
            )
            
            result = {
                'response': response,
                'is_unique': is_unique,
                'parsed_data': parsed_data,
                'suggest_processing': suggest_processing,
                'files': [],
                'diagram': None,
                'success': True
            }
            
            # Если идея валидна, создаем файлы и диаграмму
            if is_unique and (parsed_data or not is_free_form):
                data_for_files = parsed_data if parsed_data else user_data
                
                # Генерируем файлы
                try:
                    word_file, excel_file = self.file_generator.generate_files(data_for_files)
                    result['files'] = [word_file, excel_file]
                except Exception as e:
                    logging.error(f"Ошибка при создании файлов: {e}")
                
                # Генерируем диаграмму
                try:
                    diagram_file = self.diagram_generator.generate_idea_evaluation_diagram(
                        user_data, is_unique, parsed_data
                    )
                    if diagram_file:
                        result['diagram'] = diagram_file
                except Exception as e:
                    logging.error(f"Ошибка при создании диаграммы: {e}")
            
            return result
            
        except Exception as e:
            logging.error(f"Ошибка в обработке идеи: {e}")
            return {
                'response': f"⚠️ Ошибка при обработке идеи: {e}",
                'is_unique': False,
                'parsed_data': {},
                'suggest_processing': False,
                'files': [],
                'diagram': None,
                'success': False
            }

    def generate_idea_suggestions(self, user_request: str) -> str:
        """Генерация предложений идей"""
        try:
            agents_data = AgentDataProcessor.load_agents_data()
            
            existing_agents_context = ""
            if agents_data:
                agent_types = set(agent['type'] for agent in agents_data if agent['type'])
                existing_agents_context = f"Существующие типы агентов: {', '.join(agent_types)}"

            prompt = f"""
            Запрос пользователя: "{user_request}"
            {existing_agents_context}

            Сгенерируй 3-4 конкретные идеи для AI-агентов, которые могли бы помочь пользователю.
            
            Для каждой идеи предложи:
            - Название агента
            - Краткое описание (1-2 предложения)
            - Основные функции
            - Примерные преимущества
            
            Старайся предлагать разнообразные решения и избегай повторения существующих агентов.
            
            Отвечай ТОЛЬКО на русском языке, используй смайлики для наглядности.
            """
            
            logging.info(f"[GigaChat Ideas] Generating suggestions for: {user_request}")
            raw_response = get_llm().invoke(prompt)
            
            response = TextProcessor.clean_response_text(raw_response)
            return response if response else "💡 Не удалось сгенерировать идеи. Попробуйте переформулировать запрос."
            
        except Exception as e:
            logging.error(f"Ошибка при генерации идей: {e}")
            return f"⚠️ Ошибка при генерации идей: {e}"

    def get_agents_summary(self) -> Optional[str]:
        """Создание сводного файла по агентам"""
        try:
            return self.file_generator.generate_agents_summary_file("agents.xlsx")
        except Exception as e:
            logging.error(f"Ошибка при создании сводки агентов: {e}")
            return None

    def clear_user_memory(self, user_id: int) -> bool:
        """Очистка памяти пользователя"""
        return self.memory_manager.clear_user_memory(user_id)

    def get_memory_info(self, user_id: int) -> str:
        """Получение информации о памяти пользователя"""
        return self.memory_manager.get_memory_summary(user_id)

    def _handle_command(self, command: str, message: str, user_id: Optional[int], 
                       context: Optional[Dict]) -> Dict[str, Any]:
        """Обработка команд"""
        result = {}
        
        if command == 'start':
            result['response'] = self._get_start_message()
            
        elif command == 'help':
            result['response'] = self._get_help_message()
            
        elif command == 'ai_agent':
            summary_file = self.get_agents_summary()
            if summary_file:
                result['files'] = [summary_file]
                result['response'] += f"\n\n📊 Создан аналитический файл: {summary_file}"
            
        elif command == 'search_owners':
            result['response'] = self.gigachat.find_agent_owners(message)
            
        elif command == 'consultation':
            result['response'] = self.generate_idea_suggestions(message)
            
        elif command == 'idea':
            # Этот случай должен обрабатываться отдельно через process_idea
            result['response'] = "💡 Для обработки идеи используйте метод process_idea"
        
        return result

    def _get_start_message(self) -> str:
        """Приветственное сообщение"""
        return """
🤖 **Добро пожаловать в Агентолог!**

Я помогу вам с разработкой AI-агентов:

💡 **Что я умею:**
• Анализировать идеи для AI-агентов
• Проверять уникальность инициатив
• Рассчитывать стоимость разработки
• Создавать документацию и отчеты
• Искать существующих агентов и их владельцев
• Предлагать новые идеи для автоматизации

📝 **Как начать:**
• Опишите вашу идею в свободной форме
• Или заполните структурированную форму
• Спросите про существующих агентов
• Попросите помочь с идеей

🚀 **Готов помочь! Расскажите, что хотите автоматизировать?**
        """

    def _get_help_message(self) -> str:
        """Сообщение с помощью"""
        return """
🆘 **Справка по использованию Агентолога**

📋 **Основные функции:**

1️⃣ **Анализ идей:**
   • Опишите идею в свободной форме
   • Я проанализирую уникальность
   • Создам документацию и диаграмму оценки

2️⃣ **Поиск агентов:**
   • "Найди агентов для работы с документами"
   • "Кто занимается аналитикой?"

3️⃣ **Генерация идей:**
   • "Придумай идеи для HR"
   • "Что можно автоматизировать в продажах?"

4️⃣ **Расчет стоимости:**
   • Автоматически для уникальных идей
   • Детальный расчет по вопросам

5️⃣ **Получение отчетов:**
   • "Покажи всех агентов" - создам Excel с аналитикой

❓ **Проблемы:**
• Перезапустите бота
• Проверьте формат ввода
• Используйте простые фразы

💬 **Пишите как обычно, я пойму!**
        """

# Вспомогательные функции для обратной совместимости
def check_idea_with_gigachat_local(user_input: str, user_data: dict, is_free_form: bool = False) -> tuple:
    """Обратная совместимость"""
    bot = AIAgentBot()
    return bot.gigachat.check_idea_with_gigachat(user_input, user_data, is_free_form)

def check_general_message_with_gigachat(msg: str, user_id: int = None) -> tuple:
    """Обратная совместимость"""
    bot = AIAgentBot()
    return bot.gigachat.check_general_message(msg, user_id)

def find_agent_owners(query: str) -> str:
    """Обратная совместимость"""
    bot = AIAgentBot()
    return bot.gigachat.find_agent_owners(query)

def generate_idea_suggestions(user_request: str) -> str:
    """Обратная совместимость"""
    bot = AIAgentBot()
    return bot.generate_idea_suggestions(user_request)

def generate_agents_summary_file(agents_file_path: str) -> str:
    """Обратная совместимость"""
    bot = AIAgentBot()
    return bot.get_agents_summary()

def generate_files(data: dict, cost_info: str = "") -> tuple:
    """Обратная совместимость"""
    file_gen = FileGenerator()
    return file_gen.generate_files(data, cost_info)

def calculate_work_cost(parsed_data: dict, is_unique: bool = True) -> str:
    """Обратная совместимость"""
    calculator = CostCalculator()
    return calculator.calculate_work_cost(parsed_data, is_unique)

def calculate_work_cost_interactive(answers: dict, return_next: bool = False):
    """Обратная совместимость"""
    calculator = CostCalculator()
    return calculator.calculate_work_cost_interactive(answers, return_next)

def generate_cost_questions(parsed_data: dict) -> tuple:
    """Обратная совместимость"""
    cost_manager = CostCalculationManager()
    return cost_manager.generate_cost_questions(parsed_data)

def calculate_final_cost(parsed_data: dict, answers: dict, user_id: int = None) -> tuple:
    """Обратная совместимость"""
    cost_manager = CostCalculationManager()
    return cost_manager.calculate_final_cost(parsed_data, answers, user_id)

def process_cost_answers(questions: dict, user_input: str) -> tuple:
    """Обратная совместимость"""
    cost_manager = CostCalculationManager()
    return cost_manager.process_cost_answers(questions, user_input)

def handle_cost_calculation_flow(user_input: str, user_data: dict, user_id: int = None) -> tuple:
    """Обратная совместимость"""
    cost_manager = CostCalculationManager()
    return cost_manager.handle_cost_calculation_flow(user_input, user_data, user_id)

def generate_idea_evaluation_diagram(idea_data: dict, is_unique: bool, parsed_data: dict = None) -> str:
    """Обратная совместимость"""
    diagram_gen = DiagramGenerator()
    return diagram_gen.generate_idea_evaluation_diagram(idea_data, is_unique, parsed_data)

# Функции управления памятью для обратной совместимости
def add_to_memory(user_id: int, user_message: str, bot_response: str):
    """Обратная совместимость"""
    MemoryManager.add_to_memory(user_id, user_message, bot_response)

def get_conversation_context(user_id: int) -> str:
    """Обратная совместимость"""
    return MemoryManager.get_conversation_context(user_id)

def clean_response_text(text: str) -> str:
    """Обратная совместимость"""
    return TextProcessor.clean_response_text(text)

def load_agents_data() -> list:
    """Обратная совместимость"""
    return AgentDataProcessor.load_agents_data()

def safe_str(value):
    """Обратная совместимость"""
    return TextProcessor.safe_str(value)

# Основной класс для использования
__all__ = [
    'AIAgentBot',
    'GigaChatProcessor', 
    'CostCalculator',
    'FileGenerator',
    'DiagramGenerator',
    'MemoryManager',
    'TextProcessor',
    'AgentDataProcessor',
    'CostCalculationManager'
]

if __name__ == "__main__":
    # Пример использования
    bot = AIAgentBot()
    
    # Тестовое сообщение
    result = bot.process_user_message("Привет, что ты умеешь?", user_id=123)
    print("Ответ бота:", result['response'])
    
    # Тестовая идея
    test_idea = {
        "Название": "Агент для обработки заявок",
        "Что хотим улучшить?": "Автоматизировать обработку входящих заявок",
        "Масштаб процесса": "средний"
    }
    
    idea_result = bot.process_idea(test_idea, user_id=123)
    print("Анализ идеи:", idea_result['response'])
    print("Уникальна:", idea_result['is_unique'])
    print("Файлы:", idea_result['files'])
    print("Диаграмма:", idea_result['diagram'])